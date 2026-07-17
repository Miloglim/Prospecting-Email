# ponytail: one-off script — quotation docx v4 (CMA promo rates, all O/F minus USD200)
from docx import Document
from docx.shared import Pt

OUT = r"E:\Agents Basement\projects\Prospecting Email\docs\ECSA Quotation 2026-07-16 v4.docx"

HEAD = ["Hi dear friend,", "Good day!", "",
        "Pls check below rates for your reference, all DIRECT services, thanks.", ""]

# source promo 5244/5544 -200 = 5044/5344 ; Xingang 5644/6044 -200 = 5444/5844 ; MSC 3710/3910 -200 = 3510/3710
SANTOS = [
    ("SHENZHEN", "USD5044/5344", "30D", ""),
    ("NINGBO",   "USD5044/5344", "32D", ""),
    ("SHANGHAI", "USD3510/3710", "33D", "Special, subject to space"),
    ("TIANJIN",  "USD5444/5844", "41D", "via Qingdao (no direct call at Xingang)"),
]
SC = [
    ("SHENZHEN", "USD5044/5344", "32-34D", ""),
    ("NINGBO",   "USD5044/5344", "36-38D", ""),
    ("SHANGHAI", "USD5044/5344", "37-39D", ""),
    ("TIANJIN",  "USD5444/5844", "45-46D", "via Qingdao"),
]

TERMS = ["",
    "O/F = USD per 20GP / 40GP&HQ",
    "21 days free time (Shanghai-Santos special: 14 days)",
    "Valid from Jul.22 to Jul.31 (Shanghai-Santos special: subject to space)",
    "Subject to DTHC & ISPS",
    "DTHC avg: Santos ~USD340 / Navegantes ~USD215 / Itajai ~USD205 / Itapoa ~USD195 (same 20'&40')",
    "",
]
HOOKS = [
    "TIP: Itajai / Navegantes / Itapoa serve the same hinterland -- Itapoa saves you the most on DTHC if routing is flexible.",
    "40NOR PROMO: USD4000 to Santos / USD4100 to Itapoa, direct, valid Jul.21-31 -- pls advise if workable for your cargo.",
    "Urgent cargo before Jul.22? Pls advise, I'll quote the current-week level right away.",
    "",
    "Pls advise your volume and preferred carrier (if any) -- I'll lock space at the best option right away.",
    "",
    "Best regards,",
    "Zayne",
]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(2)

def para(text, bold=False):
    p = doc.add_paragraph()
    p.add_run(text).bold = bold

def rate_table(title, rows):
    para(title, bold=True)
    t = doc.add_table(rows=1 + len(rows), cols=4)
    t.style = "Table Grid"
    for i, h in enumerate(["POL", "O/F (20GP/40'&HQ)", "T/T", "REMARK"]):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
    doc.add_paragraph()

for line in HEAD:
    para(line)
rate_table("SANTOS", SANTOS)
rate_table("ITAJAI / NAVEGANTES / ITAPOA  (same rate level, all direct)", SC)
for line in TERMS:
    para(line)
for line in HOOKS:
    para(line, bold=line.startswith(("TIP:", "40NOR")))

doc.save(OUT)
print("saved:", OUT)
